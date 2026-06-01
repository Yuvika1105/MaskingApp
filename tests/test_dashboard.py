# tests/test_dashboard.py
import pytest
import io
import pandas as pd
from docx import Document
from app.data_masking.masking_policy import MaskingPolicy
from app.data_masking.masking_engine import MaskingEngine, apply_masking_strategy
from app.data_masking.file_processors import FileProcessor

@pytest.fixture
def test_policy():
    # Construct an in-memory policy or load a small dummy yaml
    return MaskingPolicy(
        column_rules={
            "Model": "MG_MODEL",
            "Dealer": "MG_BRAND",
            "Name": "PERSON",
            "Phone": "PHONE_NUMBER"
        },
        entity_rules=[
            "MG_BRAND", "MG_MODEL", "PERSON", "PHONE_NUMBER"
        ],
        replacement_map={
            "MG_BRAND": "<CLIENT>",
            "MG_MODEL": "<MODEL_LINE>",
            "PERSON": "<PERSON>",
            "PHONE_NUMBER": "<PHONE_NUMBER>"
        }
    )

@pytest.fixture
def test_engine(test_policy):
    return MaskingEngine(test_policy)

def test_apply_masking_strategies():
    # Test Redaction
    assert apply_masking_strategy("John Doe", "Redaction", "PERSON") == "<PERSON>"
    
    # Test Custom Constant
    assert apply_masking_strategy("John Doe", "Custom Constant", "PERSON", "[SECRET]") == "[SECRET]"
    
    # Test Hashing (SHA-256 starts with HASH_PERSON_)
    hash_val = apply_masking_strategy("John Doe", "Hashing", "PERSON")
    assert hash_val.startswith("[HASH_PERSON_")
    # Same value, same hash (deterministic)
    assert apply_masking_strategy("John Doe", "Hashing", "PERSON") == hash_val
    
    # Test Character Masking (standard word length > 3)
    assert apply_masking_strategy("John", "Character Masking") == "J**n"
    # Test Character Masking for email
    assert apply_masking_strategy("john.doe@gmail.com", "Character Masking") == "j******e@gmail.com"
    # Test Character Masking for phone (10 digit)
    assert apply_masking_strategy("9876543210", "Character Masking") == "987*****10"

def test_masking_engine_spans(test_engine):
    # Free-text masking
    text = "Call John Doe at 9876543210 from MG Motors showroom about HECTOR pricing."
    
    # Run with default strategies (Redaction)
    res = test_engine.mask_text(text)
    
    assert "John Doe" not in res.masked_text
    assert "9876543210" not in res.masked_text
    assert "HECTOR" not in res.masked_text
    assert "<PERSON>" in res.masked_text or "[HASH" in res.masked_text or "J**n" in res.masked_text
    assert "HR" in res.masked_text  # MG_MODEL Hector redacted abbreviation

def test_dynamic_strategies_override(test_engine):
    text = "Contact John Doe."
    
    # Apply Hashing strategy dynamically
    strategies = {"PERSON": "Hashing"}
    res = test_engine.mask_text(text, entity_strategies=strategies)
    assert "[HASH_PERSON_" in res.masked_text
    assert "John Doe" not in res.masked_text

def test_csv_file_processor(test_engine):
    # Mock CSV data
    data = {
        "Name": ["Alice", "Bob"],
        "Model": ["ASTOR", "HECTOR"],
        "Age": [25, 30]
    }
    df = pd.DataFrame(data)
    csv_bytes = df.to_csv(index=False).encode("utf-8")
    
    # Process with dynamic column settings
    column_strategies = {"Name": "Hashing", "Model": "Redaction"}
    column_constants = {"Name": ""}
    selected_columns = ["Name", "Model"]
    
    df_masked = FileProcessor.process_csv(
        file_bytes=csv_bytes,
        engine=test_engine,
        column_strategies=column_strategies,
        column_constants=column_constants,
        selected_columns=selected_columns
    )
    
    # Check that Model ASTOR is abbreviated to AR and HECTOR is abbreviated to HR
    assert df_masked.loc[0, "Model"] == "AR"
    assert df_masked.loc[1, "Model"] == "HR"
    
    # Check Name is Hashed
    assert df_masked.loc[0, "Name"].startswith("[HASH_PERSON_")
    
    # Check Age is untouched (not in selected columns)
    assert df_masked.loc[0, "Age"] == 25

def test_docx_file_processor(test_engine):
    # Create simple docx
    doc = Document()
    doc.add_paragraph("My name is John.")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    docx_bytes = doc_io.getvalue()
    
    masked_bytes = FileProcessor.process_docx(docx_bytes, test_engine)
    
    # Load back
    doc_masked = Document(io.BytesIO(masked_bytes))
    paragraph_text = doc_masked.paragraphs[0].text
    
    assert "John" not in paragraph_text
    assert "<PERSON>" in paragraph_text

def test_dynamic_custom_recognizers(test_policy):
    # Add a custom regex and keyword rule at initialization
    custom_rules = [
        {"name": "CAR_PLATE", "type": "Regex Pattern", "value": r"\b[A-Z]{2}-\d{2}\b"},
        {"name": "DESSERT", "type": "Keywords List", "value": "ice cream, chocolate cake"}
    ]
    engine = MaskingEngine(test_policy, custom_rules=custom_rules)
    
    text = "My car plate is DL-04. I would love some ice cream today."
    res = engine.mask_text(text)
    
    assert "DL-04" not in res.masked_text
    assert "<CAR_PLATE>" in res.masked_text
    
    assert "ice cream" not in res.masked_text
    assert "<DESSERT>" in res.masked_text


def test_tesla_policy_masking():
    # Load the actual Tesla policy yaml from data directory
    from config import POLICIES_DIR
    tesla_policy_path = POLICIES_DIR / "tesla_policy.yaml"
    policy = MaskingPolicy.from_yaml(str(tesla_policy_path))
    engine = MaskingEngine(policy)
    
    text = "Call Elon at 9876543210 from Tesla showroom regarding MODEL 3 delivery and code 123-AB."
    
    res = engine.mask_text(text)
    
    # Verify Tesla brand name masked as entity type in free-text
    assert "Tesla" not in res.masked_text
    assert "<TESLA_BRAND>" in res.masked_text
    
    # Verify Tesla model name abbreviated in free-text
    assert "MODEL 3" not in res.masked_text
    assert "ML" in res.masked_text  # MODEL 3 abbreviated to ML
    
    # Verify material code masked as entity type in free-text
    assert "123-AB" not in res.masked_text
    assert "<TESLA_MATERIAL>" in res.masked_text
    
    # Verify standard PII phone masked in free-text
    assert "9876543210" not in res.masked_text
    
    # Verify mask_value works correctly using replacement_map
    assert engine.mask_value("Name_1", "Tesla Motors", "Redaction") == "<CLIENT>"
    assert engine.mask_value("Material", "123-AB", "Redaction") == "<MATERIAL_CODE>"
    assert engine.mask_value("Model_Line1", "MODEL 3", "Redaction") == "ML"
