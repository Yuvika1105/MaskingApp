# app.py
import re
import io
import time
import pandas as pd
import streamlit as st
from pathlib import Path

# Import the existing generalized backend modules
from config import DEFAULT_POLICY_PATH
from app.data_masking.masking_policy import MaskingPolicy
from app.data_masking.masking_engine import MaskingEngine
from app.data_masking.file_processors import FileProcessor

# Set premium, professional page configuration
st.set_page_config(
    page_title="SafeGuard — Standalone Masking Portal",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Premium interface styling (Outfit + Plus Jakarta typography)
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&family=Plus+Jakarta+Sans:wght@300;400;500;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Plus Jakarta Sans', sans-serif;
    }
    
    .portal-title {
        font-family: 'Outfit', sans-serif;
        font-size: 2.8rem !important;
        font-weight: 800;
        background: linear-gradient(135deg, #0ea5e9, #6366f1, #a855f7);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.2rem;
        text-align: center;
    }
    
    .portal-subtitle {
        font-size: 1.15rem !important;
        color: #6b7280;
        text-align: center;
        margin-bottom: 2.5rem;
    }

    .masked-span-redact {
        background-color: #fee2e2;
        color: #dc2626;
        padding: 2px 6px;
        border-radius: 4px;
        border: 1px solid #fca5a5;
        font-weight: 600;
        font-family: monospace;
    }
    .masked-span-hash {
        background-color: #f3e8ff;
        color: #7c3aed;
        padding: 2px 6px;
        border-radius: 4px;
        border: 1px solid #d8b4fe;
        font-weight: 600;
        font-family: monospace;
    }
</style>
""", unsafe_allow_html=True)

# Main Title & Subtitle
st.markdown('<div class="portal-title">Data Masking Portal</div>', unsafe_allow_html=True)

st.divider()

# File Upload Section
st.subheader("1. Upload Document")
uploaded_file = st.file_uploader(
    "Select file to sanitize",
    type=["csv", "xlsx", "xls", "pdf", "docx", "txt", "json"],
    help="Supports CSV, Excel (.xlsx/.xls), PDF, Word (.docx), plain text (.txt), and JSON files."
)

if uploaded_file is not None:
    file_name = uploaded_file.name
    file_bytes = uploaded_file.getvalue()
    suffix = Path(file_name).suffix.lower()
    is_tabular = suffix in [".csv", ".xlsx", ".xls"]

    # Try/except block to handle file reading gracefully
    try:
        # ──────────────────────────────────────────────────────────────
        # TABULAR PATH: CSV / Excel
        # ──────────────────────────────────────────────────────────────
        if is_tabular:
            if suffix == ".csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes))

            st.success(f"Successfully loaded `{file_name}` ({len(df)} rows, {len(df.columns)} columns)")

            # Display raw data preview
            with st.expander("View Raw Data Preview (First 5 Rows)", expanded=True):
                st.dataframe(df.head(5), use_container_width=True)

            st.divider()

            # ----------------------------------------------------------
            # AUTOMATIC SCHEMA DETECTION & COMPONENT MAPPING
            # ----------------------------------------------------------
            all_column_rules = {}
            all_replacement_maps = {}
            domain_configs = {}

            try:
                p = MaskingPolicy.from_yaml(str(DEFAULT_POLICY_PATH))
                all_column_rules.update(p.column_rules)
                all_replacement_maps.update(p.replacement_map)
            except Exception:
                pass

            # Discover which specific entities correspond to the columns
            detected_entities = []
            for col in df.columns:
                entity = None
                for rule_col, rule_ent in all_column_rules.items():
                    if col.strip().lower() == rule_col.strip().lower():
                        entity = rule_ent
                        break
                if entity and entity not in detected_entities:
                    detected_entities.append(entity)

            if not detected_entities:
                detected_entities = ["PERSON", "EMAIL_ADDRESS"]

            # ----------------------------------------------------------
            # Dynamic Configuration Section
            # ----------------------------------------------------------
            st.subheader("2. Configuration & Sensitivity Rules")
            st.write("Configure column-level overrides, smart PII scanner boundaries, and custom word targets.")

            cfg_col1, cfg_col2, cfg_col3 = st.columns(3)

            with cfg_col1:
                with st.container(border=True):
                    st.write("### A: Column Masking & Abbreviation")
                    st.caption("Select columns to fully mask with `<REDACTED>` or abbreviate to first & last letters.")

                    fully_masked_cols = st.multiselect(
                        "Select columns to fully mask:",
                        options=df.columns.tolist(),
                        help="Every cell in these columns will be replaced with '<REDACTED>'."
                    )

                    remaining_opts = [c for c in df.columns if c not in fully_masked_cols]
                    abbreviated_cols = st.multiselect(
                        "Select columns to abbreviate (first & last letter):",
                        options=remaining_opts,
                        help="Every word in these columns will be abbreviated to its first and last letter (e.g. Bentley -> BY)."
                    )

            with cfg_col2:
                with st.container(border=True):
                    st.write("### B: Smart Entity Recognition")
                    st.caption("Select PII or custom entities to scan for and mask in remaining columns.")

                    entity_options = [
                        "PERSON",
                        "EMAIL_ADDRESS",
                        "PHONE_NUMBER",
                        "ORGANIZATION",
                        "CREDIT_CARD",
                        "US_SSN"
                    ]

                    pre_selected = [e for e in detected_entities if e in entity_options]
                    if not pre_selected:
                        pre_selected = ["PERSON", "EMAIL_ADDRESS"]

                    selected_entities = st.multiselect(
                        "Select entities to detect and mask:",
                        options=entity_options,
                        default=pre_selected,
                        help="The engine will scan cell text in remaining columns and mask any matching instances."
                    )

            with cfg_col3:
                with st.container(border=True):
                    st.write("### C: Custom Term Masking")
                    st.caption("Provide a comma-separated list of specific words or phrases (e.g. `MG, Tata, Tesla`) to always mask.")

                    custom_redact_terms = st.text_input(
                        "Enter custom terms to mask (comma-separated):",
                        placeholder="e.g., MG, Tata, Tesla",
                        help="These words will be case-insensitively found and replaced with '#' everywhere in the file."
                    )

                    custom_redact_list = [t.strip() for t in custom_redact_terms.split(",") if t.strip()] if custom_redact_terms else []
                    if custom_redact_list:
                        st.info(f"Targeting {len(custom_redact_list)} custom terms for masking.")

            st.divider()

            # Execution & Processing Section
            st.subheader("3. Execution & Sanitization Engine")

            if st.button("Apply Security Policies & Run Masking", type="primary", use_container_width=True):
                with st.spinner("Executing dynamic data sanitization... Please wait."):
                    start_time = time.time()

                    sanitized_df = df.copy()

                    base_policy = MaskingPolicy(
                        column_rules={c: all_column_rules[c] for c in df.columns if c in all_column_rules},
                        entity_rules=selected_entities,
                        replacement_map=all_replacement_maps,
                        domain_config=domain_configs
                    )

                    engine = MaskingEngine(base_policy)

                    def abbreviate_text(text):
                        if not text or pd.isna(text):
                            return ""
                        text_str = str(text).strip()
                        if not text_str or text_str.lower() in ["nan", "nat", "<na>", "none"]:
                            return ""

                        def abbrev_word(w):
                            w_clean = w.strip()
                            if not w_clean.isalnum():
                                return w_clean
                            if len(w_clean) <= 1:
                                return w_clean.upper()
                            return (w_clean[0] + w_clean[-1]).upper()

                        words = text_str.split()
                        return " ".join(abbrev_word(w) for w in words)

                    def apply_custom_redactions(val):
                        if pd.isna(val):
                            return val
                        val_str = str(val).strip()
                        if not val_str or val_str.lower() in ["nan", "nat", "<na>", "none"]:
                            return ""
                        masked_val = val_str
                        if custom_redact_list:
                            for term in custom_redact_list:
                                tag_pattern = re.compile(rf"<[A-Z_]*{re.escape(term.upper())}[A-Z_]*>", re.IGNORECASE)
                                masked_val = tag_pattern.sub("#", masked_val)
                                substring_pattern = re.compile(re.escape(term), re.IGNORECASE)
                                masked_val = substring_pattern.sub("#", masked_val)
                        return masked_val

                    for col in fully_masked_cols:
                        sanitized_df[col] = "<REDACTED>"

                    for col in abbreviated_cols:
                        sanitized_df[col] = sanitized_df[col].apply(apply_custom_redactions).apply(abbreviate_text)

                    remaining_cols = [c for c in df.columns if c not in fully_masked_cols and c not in abbreviated_cols]

                    def mask_cell(val):
                        if pd.isna(val):
                            return val
                        val_str = str(val).strip()
                        if not val_str or val_str.lower() in ["nan", "nat", "<na>", "none"]:
                            return ""
                        val_redacted = apply_custom_redactions(val_str)
                        res = engine.mask_text(val_redacted)
                        masked_val = res.masked_text
                        masked_val = apply_custom_redactions(masked_val)
                        return masked_val

                    for col in remaining_cols:
                        sanitized_df[col] = sanitized_df[col].apply(mask_cell)

                    elapsed_time = time.time() - start_time

                    st.success(f"Processing Completed! Masked successfully in {elapsed_time:.3f} seconds.")

                    st.divider()

                    # Review & Export Section
                    st.subheader("4. Review & Secure Export")

                    with st.expander("View Masked Data Preview (First 5 Rows)", expanded=True):
                        st.dataframe(sanitized_df.head(5), use_container_width=True)

                    exp_c1, exp_c2 = st.columns(2)

                    csv_buffer = io.StringIO()
                    sanitized_df.to_csv(csv_buffer, index=False)
                    csv_bytes = csv_buffer.getvalue().encode("utf-8")
                    exp_c1.download_button(
                        label="📥 Download as CSV",
                        data=csv_bytes,
                        file_name=f"masked_{Path(file_name).stem}.csv",
                        mime="text/csv",
                        use_container_width=True
                    )

                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
                        sanitized_df.to_excel(writer, index=False)
                    exp_c2.download_button(
                        label="📥 Download as Excel",
                        data=excel_buffer.getvalue(),
                        file_name=f"masked_{Path(file_name).stem}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )

        # ──────────────────────────────────────────────────────────────
        # UNSTRUCTURED PATH: PDF / DOCX / TXT / JSON
        # ──────────────────────────────────────────────────────────────
        else:
            st.success(f"Successfully loaded `{file_name}` ({len(file_bytes):,} bytes)")

            # Load engine with default policy
            policy = MaskingPolicy.from_yaml(str(DEFAULT_POLICY_PATH))
            engine = MaskingEngine(policy)

            # Extract raw text for preview
            if suffix == ".txt":
                raw_text = file_bytes.decode("utf-8", errors="ignore")
            elif suffix == ".json":
                import json
                raw_text = file_bytes.decode("utf-8", errors="ignore")
                try:
                    raw_text = json.dumps(json.loads(raw_text), indent=2)
                except Exception:
                    pass
            elif suffix == ".pdf":
                import pdfplumber
                extracted = []
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            extracted.append(t)
                raw_text = "\n".join(extracted)
            elif suffix == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(file_bytes))
                raw_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                raw_text = file_bytes.decode("utf-8", errors="ignore")

            # Show raw preview
            with st.expander("View Raw Document Content (first 500 chars)", expanded=True):
                st.text(raw_text[:500] + ("..." if len(raw_text) > 500 else ""))

            st.divider()

            # Configuration
            st.subheader("2. Configuration & Sensitivity Rules")

            cfg_col1, cfg_col2 = st.columns(2)

            with cfg_col1:
                with st.container(border=True):
                    st.write("### A: Smart Entity Recognition")
                    st.caption("Choose which PII entity types to detect and redact.")
                    entity_options = ["PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "ORGANIZATION", "CREDIT_CARD", "US_SSN"]
                    selected_entities = st.multiselect(
                        "Entities to detect:",
                        options=entity_options,
                        default=["PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER"],
                    )

            with cfg_col2:
                with st.container(border=True):
                    st.write("### B: Custom Term Redaction")
                    st.caption("Comma-separated list of custom words/phrases to always mask.")
                    custom_terms_input = st.text_input(
                        "Custom terms (comma-separated):",
                        placeholder="e.g., Acme Corp, Project X"
                    )
                    custom_words = [t.strip() for t in custom_terms_input.split(",") if t.strip()] if custom_terms_input else []
                    if custom_words:
                        st.info(f"Targeting {len(custom_words)} custom terms.")

            st.divider()
            st.subheader("3. Execution & Sanitization Engine")

            if st.button("Apply Security Policies & Run Masking", type="primary", use_container_width=True):
                with st.spinner("Scanning and redacting document..."):
                    start_time = time.time()

                    # Update policy entity rules to user selection
                    engine.policy.entity_rules = selected_entities
                    engine.presidio_entities = selected_entities

                    # Process by file type
                    if suffix == ".pdf":
                        masked_text = FileProcessor.process_pdf(file_bytes, engine)
                    elif suffix == ".docx":
                        masked_docx_bytes = FileProcessor.process_docx(file_bytes, engine)
                        from docx import Document as DocxDocument
                        doc_masked = DocxDocument(io.BytesIO(masked_docx_bytes))
                        masked_text = "\n".join([p.text for p in doc_masked.paragraphs if p.text.strip()])
                    elif suffix == ".json":
                        masked_text = FileProcessor.process_json(file_bytes, engine)
                    else:
                        masked_text = FileProcessor.process_txt(file_bytes, engine)

                    # Apply custom word redactions on top
                    if custom_words:
                        for word in custom_words:
                            pattern = re.compile(rf"\b{re.escape(word)}\b", re.IGNORECASE)
                            masked_text = pattern.sub("[REDACTED]", masked_text)

                    elapsed_time = time.time() - start_time
                    st.success(f"Processing Completed in {elapsed_time:.3f} seconds.")

                    st.divider()
                    st.subheader("4. Review & Secure Export")

                    # Side-by-side preview
                    left_col, right_col = st.columns(2)

                    with left_col:
                        st.markdown("#### 📥 Original Document")
                        st.markdown(
                            f'<div style="height:380px;overflow-y:auto;padding:1rem;background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;font-size:0.9rem;white-space:pre-wrap;">{raw_text[:3000]}</div>',
                            unsafe_allow_html=True
                        )

                    with right_col:
                        st.markdown("#### 📤 Masked Document")
                        # Color-code redaction tokens
                        highlighted = re.sub(
                            r"(<([A-Z_]+)>)",
                            r'<span class="masked-span-redact">\1</span>',
                            masked_text[:3000].replace("\n", "<br>")
                        )
                        highlighted = re.sub(
                            r"(\[HASH_([A-Z_]+)_([A-Z0-9]+)\])",
                            r'<span class="masked-span-hash">\1</span>',
                            highlighted
                        )
                        st.markdown(
                            f'<div style="height:380px;overflow-y:auto;padding:1rem;background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;font-size:0.9rem;">{highlighted}</div>',
                            unsafe_allow_html=True
                        )

                    st.markdown("---")

                    exp_c1, exp_c2 = st.columns(2)

                    # TXT download (always available)
                    exp_c1.download_button(
                        label="📥 Download Masked Text (.txt)",
                        data=masked_text.encode("utf-8"),
                        file_name=f"masked_{Path(file_name).stem}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )

                    # DOCX download (when original was .docx, re-export masked docx)
                    if suffix == ".docx":
                        exp_c2.download_button(
                            label="📥 Download Masked Word Doc (.docx)",
                            data=masked_docx_bytes,
                            file_name=f"masked_{Path(file_name).stem}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

    except Exception as err:
        st.error(f"An error occurred while reading or processing the file: {str(err)}")
        st.exception(err)

else:
    st.info("Upload any file above — CSV, Excel, PDF, Word (.docx), plain text (.txt), or JSON to start masking.")
