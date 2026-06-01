import json
import io
import pandas as pd
import pdfplumber
from docx import Document
from pathlib import Path
from typing import Dict, Optional, Union

from app.data_masking.masking_engine import MaskingEngine

class FileProcessor:
    @staticmethod
    def process_txt(
        content: Union[str, bytes], 
        engine: MaskingEngine, 
        strategies: Optional[Dict[str, str]] = None,
        constants: Optional[Dict[str, str]] = None
    ) -> str:
        text_str = content if isinstance(content, str) else content.decode("utf-8", errors="ignore")
        return engine.mask_text(text_str, strategies, constants).masked_text

    @staticmethod
    def process_pdf(
        file_bytes: bytes, 
        engine: MaskingEngine,
        strategies: Optional[Dict[str, str]] = None,
        constants: Optional[Dict[str, str]] = None
    ) -> str:
        extracted = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text: 
                    extracted.append(text)
        full_text = "\n".join(extracted)
        return engine.mask_text(full_text, strategies, constants).masked_text

    @staticmethod
    def process_docx(
        file_bytes: bytes, 
        engine: MaskingEngine,
        strategies: Optional[Dict[str, str]] = None,
        constants: Optional[Dict[str, str]] = None
    ) -> bytes:
        # Load the binary docx file and replace paragraph and table texts in place, returning binary data
        doc = Document(io.BytesIO(file_bytes))
        
        # Mask Paragraphs
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                # We replace run text carefully or the whole paragraph to preserve formatting as much as possible
                # Simple replacement of paragraph text is safest, though it can reset styles on direct runs.
                # A more sophisticated way is to mask the whole text, and if runs exist, distribute the replacement
                # or just set the paragraph text directly. Let's do direct text replacement since it's robust.
                p.text = engine.mask_text(p.text, strategies, constants).masked_text
                
        # Mask Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text and p.text.strip():
                            p.text = engine.mask_text(p.text, strategies, constants).masked_text
                            
        out_stream = io.BytesIO()
        doc.save(out_stream)
        return out_stream.getvalue()

    @staticmethod
    def process_csv(
        file_bytes: bytes, 
        engine: MaskingEngine,
        column_strategies: Dict[str, str],
        column_constants: Dict[str, str],
        selected_columns: list,
        row_filter_col: Optional[str] = None,
        row_filter_op: Optional[str] = None,
        row_filter_val: Optional[str] = None
    ) -> pd.DataFrame:
        df = pd.read_csv(io.BytesIO(file_bytes))
        
        # Apply Row Filtering Condition if configured
        rows_to_mask = pd.Series(True, index=df.index)
        if row_filter_col and row_filter_op and row_filter_val is not None:
            if row_filter_col in df.columns:
                col_data = df[row_filter_col].astype(str)
                if row_filter_op == "Equals":
                    rows_to_mask = col_data == str(row_filter_val)
                elif row_filter_op == "Contains":
                    rows_to_mask = col_data.str.contains(str(row_filter_val), case=False, na=False)
                elif row_filter_op == "Greater than":
                    try:
                        rows_to_mask = pd.to_numeric(df[row_filter_col]) > float(row_filter_val)
                    except ValueError:
                        pass
                elif row_filter_op == "Less than":
                    try:
                        rows_to_mask = pd.to_numeric(df[row_filter_col]) < float(row_filter_val)
                    except ValueError:
                        pass
        
        # Apply Column Strategies
        for col in df.columns:
            if col in selected_columns:
                strat = column_strategies.get(col, "Redaction")
                const_val = column_constants.get(col, None)
                
                # Apply mask to cells in rows matching filter
                def mask_cell(v, index):
                    if not rows_to_mask.loc[index]:
                        return v
                    return engine.mask_value(col, str(v), strat, const_val)

                # Vectorized or apply-based cell-level replacement
                df[col] = df.apply(lambda r: mask_cell(r[col], r.name), axis=1)
                
        return df

    @staticmethod
    def process_excel(
        file_bytes: bytes, 
        engine: MaskingEngine,
        column_strategies: Dict[str, str],
        column_constants: Dict[str, str],
        selected_columns: list,
        row_filter_col: Optional[str] = None,
        row_filter_op: Optional[str] = None,
        row_filter_val: Optional[str] = None
    ) -> pd.DataFrame:
        df = pd.read_excel(io.BytesIO(file_bytes))
        
        # Apply Row Filtering Condition if configured
        rows_to_mask = pd.Series(True, index=df.index)
        if row_filter_col and row_filter_op and row_filter_val is not None:
            if row_filter_col in df.columns:
                col_data = df[row_filter_col].astype(str)
                if row_filter_op == "Equals":
                    rows_to_mask = col_data == str(row_filter_val)
                elif row_filter_op == "Contains":
                    rows_to_mask = col_data.str.contains(str(row_filter_val), case=False, na=False)
                elif row_filter_op == "Greater than":
                    try:
                        rows_to_mask = pd.to_numeric(df[row_filter_col]) > float(row_filter_val)
                    except ValueError:
                        pass
                elif row_filter_op == "Less than":
                    try:
                        rows_to_mask = pd.to_numeric(df[row_filter_col]) < float(row_filter_val)
                    except ValueError:
                        pass
        
        # Apply Column Strategies
        for col in df.columns:
            if col in selected_columns:
                strat = column_strategies.get(col, "Redaction")
                const_val = column_constants.get(col, None)
                
                # Apply mask to cells in rows matching filter
                def mask_cell(v, index):
                    if not rows_to_mask.loc[index]:
                        return v
                    return engine.mask_value(col, str(v), strat, const_val)
                
                df[col] = df.apply(lambda r: mask_cell(r[col], r.name), axis=1)
                
        return df

    @staticmethod
    def process_json(
        file_bytes: bytes, 
        engine: MaskingEngine,
        strategies: Optional[Dict[str, str]] = None,
        constants: Optional[Dict[str, str]] = None
    ) -> str:
        # We can parse it as text, mask it to cover all occurrences, and return pretty-printed JSON text
        raw_text = file_bytes.decode("utf-8", errors="ignore")
        try:
            parsed = json.loads(raw_text)
            # Traverse and mask values to preserve valid JSON structure
            def recursive_mask(obj):
                if isinstance(obj, dict):
                    return {k: recursive_mask(v) for k, v in obj.items()}
                elif isinstance(obj, list):
                    return [recursive_mask(elem) for elem in obj]
                elif isinstance(obj, str):
                    return engine.mask_text(obj, strategies, constants).masked_text
                return obj
            
            masked_obj = recursive_mask(parsed)
            return json.dumps(masked_obj, indent=2)
        except json.JSONDecodeError:
            # Fallback to direct text masking if invalid JSON
            return engine.mask_text(raw_text, strategies, constants).masked_text
